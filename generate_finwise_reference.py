from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/velkumaran/Documents/Codex/2026-04-28/files-mentioned-by-the-user-app")
OUTPUT_DOCX = ROOT / "Finwise-Architecture-and-Deployment-Guide.docx"
ARCH_IMG = ROOT / "finwise-architecture-diagram.png"
FLOW_IMG = ROOT / "finwise-delivery-flow.png"

ACCENT = RGBColor(160, 120, 64)
TEXT = RGBColor(28, 26, 23)
MUTED = RGBColor(107, 101, 96)
BG = (248, 244, 238)
CARD = (255, 255, 255)
CARD_ALT = (245, 239, 230)
BORDER = (217, 206, 190)
GREEN = (5, 150, 105)
RED = (220, 38, 38)
BLUE = (79, 70, 229)
PURPLE = (129, 140, 248)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DDD5CB"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def style_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = TEXT

    for name, size in [("Heading 1", 18), ("Heading 2", 13.5), ("Heading 3", 11.5)]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = ACCENT if name == "Heading 1" else TEXT


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(10)
    run = p.add_run("FINWISE")
    run.font.name = "Aptos Display"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Architecture, Build, AI, Auth, and Deployment Guide")
    run.font.name = "Aptos Display"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run(
        "A complete reference for how the Finwise finance tracker was designed, improved, and deployed from scratch."
    )
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.autofit = False
    widths = [Inches(1.9), Inches(4.9)]
    for row in meta.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
    items = [
        ("Prepared for", "Vel Kumaran"),
        ("Project", "Finwise: Premium AI Finance Tracker"),
        ("Scope", "Product architecture, implementation journey, environment setup, troubleshooting, and production deployment"),
        ("Generated on", datetime.now().strftime("%d %B %Y, %I:%M %p")),
    ]
    for r, (k, v) in enumerate(items):
        left, right = meta.rows[r].cells
        set_cell_shading(left, "F3EBDD")
        left.paragraphs[0].add_run(k).bold = True
        right.paragraphs[0].add_run(v)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "This guide documents the real implementation choices used in the live application, including Firebase authentication, Firestore transaction storage, Firebase-backed profile persistence, AI and market proxy architecture, and Vercel serverless deployment."
    )
    run.italic = True
    run.font.color.rgb = MUTED
    doc.add_page_break()


def add_toc(doc: Document):
    doc.add_heading("Contents", level=1)
    items = [
        "1. Product Overview",
        "2. Final Feature Set",
        "3. Technology Stack",
        "4. High-Level Architecture",
        "5. Frontend Application Design",
        "6. Data Model and Core Calculations",
        "7. Authentication System",
        "8. AI Insights Platform",
        "9. Market Holdings and Valuation Platform",
        "10. Local Development Workflow",
        "11. Production Deployment on Vercel",
        "12. Environment Variables and Secrets",
        "13. Security and Reliability Decisions",
        "14. Performance Optimization and Code Splitting",
        "15. Implementation Journey and Major Decisions",
        "16. Testing and Validation Checklist",
        "17. Troubleshooting Reference",
        "18. File and Folder Reference",
        "19. Recommended Next Steps",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)
    doc.add_page_break()


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def draw_box(draw, xy, text, fill, outline, title_fill=None, title=None, font=None, small=None):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    tx = x1 + 18
    ty = y1 + 18
    if title:
      title_h = 34
      draw.rounded_rectangle((x1 + 14, y1 + 12, x2 - 14, y1 + 12 + title_h), radius=12, fill=title_fill or fill)
      draw.text((x1 + 26, y1 + 18), title, font=font, fill=(25, 25, 25))
      ty = y1 + 58
    for line in text.split("\n"):
        draw.text((tx, ty), line, font=small or font, fill=(30, 30, 30))
        ty += (small or font).size + 10


def arrow(draw, start, end, fill=(90, 90, 90), width=4):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex > sx else -1
        draw.polygon([(ex, ey), (ex - 16 * direction, ey - 8), (ex - 16 * direction, ey + 8)], fill=fill)
    else:
        direction = 1 if ey > sy else -1
        draw.polygon([(ex, ey), (ex - 8, ey - 16 * direction), (ex + 8, ey - 16 * direction)], fill=fill)


def create_diagrams():
    font_title = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 28)
    font_box = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 18)
    font_small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 15)

    img = Image.new("RGB", (1800, 1060), BG)
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "Finwise System Architecture", font=font_title, fill=(35, 30, 24))

    draw_box(draw, (60, 120, 470, 360), "React + Vite SPA\nModular finance UI\nLazy-loaded feature views\nTheme system + charts", CARD, BORDER, title="Client App", title_fill=(243, 235, 221), font=font_box, small=font_small)
    draw_box(draw, (560, 120, 980, 360), "Email/password + magic link\nGoogle sign-in\nVerification + reset + email change", CARD, BORDER, title="Firebase Auth", title_fill=(231, 244, 236), font=font_box, small=font_small)
    draw_box(draw, (1060, 120, 1490, 360), "users/{uid}/expenses\nusers/{uid}/meta/appState\nTransactions + profile state in Firestore\nDerived KPIs in app memory", CARD, BORDER, title="Data Layer", title_fill=(237, 240, 255), font=font_box, small=font_small)
    draw_box(draw, (60, 470, 470, 760), "Dashboard\nAI Insights\nAnalytics & Reports\nGoals / Net Worth / Import / Settings", CARD, BORDER, title="Feature Modules", title_fill=(243, 235, 221), font=font_box, small=font_small)
    draw_box(draw, (560, 470, 980, 760), "Local proxy: server/ai-proxy.mjs\nShared runtimes: ai-runtime + market-runtime\n/api/ai/* and /api/market/* Vercel functions\nProvider routing + health status", CARD, BORDER, title="Backend Runtime", title_fill=(231, 244, 236), font=font_box, small=font_small)
    draw_box(draw, (1060, 470, 1490, 760), "OpenAI / OpenRouter / Anthropic\nAlpha Vantage stock API\nAMFI-backed mutual fund NAV feed\nBrowser never stores secret keys", CARD, BORDER, title="External Providers", title_fill=(237, 240, 255), font=font_box, small=font_small)
    draw_box(draw, (560, 850, 1240, 1000), "Deployment target: Vercel\nFrontend served as static app\n/api handlers deployed as serverless functions\nEnvironment variables stored in Vercel project settings", CARD_ALT, BORDER, title="Production Runtime", title_fill=(243, 235, 221), font=font_box, small=font_small)

    arrow(draw, (470, 210), (560, 210))
    arrow(draw, (980, 210), (1060, 210))
    arrow(draw, (265, 360), (265, 470))
    arrow(draw, (770, 360), (770, 470))
    arrow(draw, (1275, 360), (1275, 470))
    arrow(draw, (470, 615), (560, 615))
    arrow(draw, (980, 615), (1060, 615))
    arrow(draw, (900, 760), (900, 850))

    img.save(ARCH_IMG)

    img2 = Image.new("RGB", (1800, 980), BG)
    draw = ImageDraw.Draw(img2)
    draw.text((60, 40), "Implementation and Delivery Lifecycle", font=font_title, fill=(35, 30, 24))
    steps = [
        ("1. Core Template", "Existing React finance tracker used as the design and state-management base."),
        ("2. Product Expansion", "Markets removed. AI, analytics, goals, recurring transactions, smart alerts, and net worth added."),
        ("3. App Refactor", "Large App.jsx split into modular components, shared utilities, and centralized styles."),
        ("4. AI Security", "Client-side keys removed. Local proxy and Vercel serverless API endpoints introduced."),
        ("5. Auth Rework", "Apple and phone OTP removed. Email link, verification, password reset, and email change added."),
        ("6. Market Holdings", "Portfolio holdings, official NSE EOD support for Indian symbols, global equities, crypto, commodities, and currency-aware valuation were added."),
        ("7. Persistence Upgrade", "Goals, assets, liabilities, holdings, snapshots, preferences, and vault metadata moved into Firebase-backed profile state."),
        ("8. UX Polish", "AI workspace refined, onboarding helpers added, goals/net worth upgraded, imports improved, and multilingual labels expanded."),
        ("9. Phase 2 Planning", "Demo mode, guided first-month flow, FIRE/EMI/insurance planners, portfolio intelligence, documents vault, and timeline views were added."),
        ("10. Auth + Admin", "Passkey quick sign-in, mobile login polish, role-based admin controls, and backend-aware admin gating were added."),
        ("11. PWA + Vault Sync", "Manifest, service worker, install surface, local-first vault uploads, and cloud sync status handling were introduced."),
        ("12. Deployment", "Git hygiene fixed, Vercel env vars configured, serverless handlers validated, and production PWA behavior prepared."),
    ]
    y = 120
    for idx, (title, body) in enumerate(steps):
        fill = CARD if idx % 2 == 0 else CARD_ALT
        draw_box(draw, (100, y, 1700, y + 86), title + "\n" + body, fill, BORDER, font=font_box, small=font_small)
        if idx < len(steps) - 1:
            arrow(draw, (900, y + 86), (900, y + 116))
        y += 110
    img2.save(FLOW_IMG)


def add_image_center(doc: Document, path: Path, width_inches: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr[i], "F3EBDD")
        set_cell_border(hdr[i])
        if widths:
            hdr[i].width = widths[i]
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(row[i])
            if widths:
                row[i].width = widths[i]
    return table


def build_document():
    create_diagrams()
    doc = Document()
    style_doc(doc)
    add_page_number(doc.sections[0])
    add_cover(doc)
    add_toc(doc)

    doc.add_heading("1. Product Overview", level=1)
    add_paragraph(doc, "Finwise is a premium personal finance tracker built as a React single-page application with Firebase for authentication, Firestore for account state, Firebase Storage for vault files, and backend AI and market-data runtimes for secure provider access. The app evolved from a lighter expense tracker into a broader command center covering income, expenses, investments, insurance, live-valued market holdings, planners, goals, net worth, analytics, AI reasoning, import/export, documents, timeline management, and role-aware administration.")
    add_bullets(doc, [
        "The final product is opinionated around practical day-to-day financial management rather than stock-market watching.",
        "The app is designed to support both manual entry and CSV import, then turn that data into diagnostics, alerts, planning, portfolio valuation, and AI-assisted explanation.",
        "The production target is Vercel, with Firebase as the identity/data layer, Firebase Storage as the vault file layer, and AI plus market providers behind serverless endpoints.",
        "The deployed app is now prepared as a PWA with service-worker caching, manifest metadata, and install prompts on supported browsers.",
    ])

    doc.add_heading("2. Final Feature Set", level=1)
    add_table(
        doc,
        ["Area", "Final capability"],
        [
            ("Overview Dashboard", "Cash position, savings rate, net worth, budget progress, runway, investment rate, goal funding gap, charts, alerts, onboarding setup guidance"),
            ("Transactions", "Income, expense, investment, and insurance types with recurring transaction support"),
            ("AI Insights", "Heuristic analysis, provider-backed insights, ask-anything financial Q&A, anomaly view, savings and investment suggestions"),
            ("Analytics", "Twelve-month trend charts, year-over-year comparisons, average metrics, category trend lines, spending heat map, CSV and PDF export"),
            ("Goals", "Goal templates, progress tracking, monthly pace estimation, status flags, contribution actions, emergency fund planner, EMI planner, retirement planner, FIRE calculator, insurance adequacy, and cash allocation planning"),
            ("Net Worth", "Manual assets, liabilities, liquidity buffer, debt ratio, balance-sheet diagnostics, starter templates, live market holdings, XIRR/CAGR-style estimates, sector split, country exposure, SIP signals, and rebalance hints"),
            ("Timeline", "Calendar-like timeline for recurring bills, goal milestones, vault reminders, and planner rhythms"),
            ("Vault", "Document metadata, local-first uploads, cloud sync, renewals, reminders, export/import, and migration of old local files"),
            ("Import", "CSV upload, column mapping, preview, validation, staged import"),
            ("Settings and Account", "Theme, language, budget, notifications, install app, account controls, passkey management, and admin-only backend/provider configuration"),
        ],
        widths=[Inches(1.7), Inches(5.6)],
    )

    doc.add_heading("3. Technology Stack", level=1)
    add_table(
        doc,
        ["Layer", "Choice", "Reason"],
        [
            ("Frontend", "React 19 + Vite 8", "Fast development loop, simple bundling, strong compatibility with modular SPA structure"),
            ("UI + Visuals", "Custom CSS + Recharts", "Allowed a tailored product style while still supporting premium charts and dashboards"),
            ("Authentication", "Firebase Auth", "Rapid support for email/password, Google sign-in, password reset, email verification, and magic-link sign-in"),
            ("Data Storage", "Cloud Firestore + Firebase Storage", "Per-user transaction/profile state in Firestore plus vault file storage in Firebase Storage"),
            ("Backend Runtime", "Node.js server + Vercel serverless functions", "Kept AI and market-data keys away from the browser and unified runtime behavior across local and deployed environments"),
            ("Deployment", "Vercel + PWA shell", "Static frontend hosting, `/api` serverless functions, service-worker delivery, and environment variable management"),
        ],
        widths=[Inches(1.2), Inches(1.9), Inches(4.2)],
    )

    doc.add_heading("4. High-Level Architecture", level=1)
    add_paragraph(doc, "The architecture has five major zones: the client app, Firebase-managed identity and persistence, Firebase Storage for vault documents, the backend AI and market runtimes, and the external managed providers. Firebase handles identity, transactions, role-aware profile state, and document metadata. Firebase Storage holds uploaded vault files. The backend runtimes handle provider selection, admin-aware health reporting, market refresh, FX, and key isolation. The Vercel deployment serves both the frontend bundle and the production API functions, plus the PWA shell.")
    add_image_center(doc, ARCH_IMG, 6.8)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Figure 1. Finwise system architecture")
    run.italic = True
    run.font.color.rgb = MUTED

    doc.add_heading("5. Frontend Application Design", level=1)
    add_paragraph(doc, "The main stateful shell lives in `src/App.jsx`. It owns authentication state, transaction subscription, Firebase-backed profile state, derived metrics, role detection, tab routing, AI request orchestration, market-holdings persistence, vault document state, install prompt state, budget settings, goals, assets, liabilities, and notification toasts. Feature screens are modular React components loaded on demand through lazy loading and a `Suspense` boundary.")
    add_bullets(doc, [
        "The app uses one navigation model for desktop sidebar and one for mobile bottom navigation, both driven by the same `activeTab` state.",
        "Feature modules include Dashboard, AIInsights, AnalyticsReports, GoalsTargets, NetWorthTracker, AddForm, History, ImportPage, DocumentsVault, TimelineView, LoginPage, and Settings.",
        "Shared logic lives in `src/lib/utils.js`, `src/lib/ai.js`, `src/lib/market.js`, `src/lib/passkey.js`, `src/lib/i18n.js`, and `src/lib/vaultStorage.js`.",
        "The CSS is centralized in `src/styles/appStyles.js`, giving the product a consistent visual system across cards, charts, forms, empty states, and premium onboarding surfaces.",
    ])

    doc.add_heading("6. Data Model and Core Calculations", level=1)
    add_paragraph(doc, "Transactions are stored in Firestore under `users/{uid}/expenses`. Each record contains amount, type, category, note, date, recurring, and recurringFrequency. User-owned planning and preference state is stored in Firestore under `users/{uid}/meta/appState`, including goals, manual assets, liabilities, market holdings, portfolio snapshots, AI preferences, budget, language, onboarding state, planner state, vault document metadata, and notification toggles. Vault files themselves live in Firebase Storage, while local IndexedDB acts as a staging layer during local-first uploads. Local storage remains only as a migration fallback and for a few transient sign-in helpers.")
    add_table(
        doc,
        ["Metric", "How it is calculated"],
        [
            ("Balance", "income - expense - investment - insurance for the selected period"),
            ("Savings Rate", "((income - expense) / income) * 100"),
            ("Tracked Cash", "allTimeIncome - allTimeExpense - allTimeInvestment - allTimeInsurance"),
            ("Tracked Investments", "nonMarketInvestment + liveHoldingsValue when holdings exist, otherwise cash-ledger market investment total"),
            ("Tracked Net Worth", "trackedCash + trackedInvestments + manualAssetTotal - liabilityTotal, with FX-aware conversion when global holdings are shown in INR or BOTH views"),
            ("Recurring Outflow", "sum of selected-period recurring non-income transactions"),
            ("Budget Progress", "selected-period expense / budget"),
            ("Unusual Transaction", "expense above 1.8x its category historical average when enough history exists"),
        ],
        widths=[Inches(1.8), Inches(5.5)],
    )

    doc.add_heading("7. Authentication System", level=1)
    add_paragraph(doc, "Authentication began as a broader experiment with email/password, Google, Apple, and phone OTP. In practice, the stable and user-friendly production mix became email/password, Google sign-in, email-link magic-link authentication, and a passkey quick-sign-in helper for supported secure environments. Apple sign-in and phone OTP were removed from the experience because they added setup and deliverability complexity without improving reliability for this product.")
    add_bullets(doc, [
        "Email/password supports account creation and sign-in.",
        "Google OAuth supports quick sign-in with popup or redirect depending on screen size.",
        "Email link sign-in acts as a passwordless fallback and is completed with `isSignInWithEmailLink` and `signInWithEmailLink`.",
        "Passkeys are treated as a premium quick-unlock layer on secure deployed origins rather than as a fully separate identity system.",
        "Settings exposes email verification, password reset, and verified email change via Firebase APIs.",
        "Admin access is role-based: the same auth methods are used, but admin-only controls are unlocked via Firebase custom claims or backend email allowlisting.",
    ])
    add_paragraph(doc, "Important implementation note: the app stores the pending email-link address in local storage, sends the link with `sendSignInLinkToEmail`, and on return completes the flow by checking the current URL for a sign-in link. If the saved email is missing, the user is prompted to re-enter the same email address for completion.")

    doc.add_heading("8. AI Insights Platform", level=1)
    add_paragraph(doc, "The AI layer was designed in two tiers. Tier one is a local heuristic analysis engine inside `src/lib/ai.js`, which can always generate structured advice even when no provider key is available. Tier two is the provider-backed AI runtime, which sends structured financial context to OpenAI, OpenRouter, or Anthropic through the backend proxy. The AI configuration surface is now admin-aware, so provider health and model configuration are not shown to ordinary users.")
    add_table(
        doc,
        ["Part", "Responsibility"],
        [
            ("`src/lib/ai.js`", "Build heuristic report, ask-answer fallback logic, backend health fetch, API client requests"),
            ("`server/ai-runtime.mjs`", "Load local `.env`, report backend health, build prompts, route provider calls, format JSON responses, enforce admin-aware health details"),
            ("`server/ai-proxy.mjs`", "Local Node HTTP server for development on port 8787"),
            ("`api/ai/*.js`", "Vercel production handlers for health, insights, and chat query endpoints"),
        ],
        widths=[Inches(2.0), Inches(5.3)],
    )
    add_paragraph(doc, "The browser never stores provider keys. The frontend only chooses provider and model name. Secret keys live in the proxy environment, both locally and in Vercel. Admin-only health surfaces are intentionally separated from the normal user settings experience.")

    doc.add_heading("9. Market Holdings and Valuation Platform", level=1)
    add_paragraph(doc, "Market holdings were added as a real portfolio layer rather than being treated as plain cashflow categories. Users can record stocks, mutual funds, crypto, and commodities separately from transaction entries, then refresh values manually once per day. Over time the valuation platform evolved into an India-first plus global fallback model: official NSE end-of-day pricing for Indian symbols, provider-backed paths for global equities and crypto, AMFI-backed mutual fund NAV, and currency-aware presentation for INR/USD/BOTH views.")
    add_table(
        doc,
        ["Part", "Responsibility"],
        [
            ("`src/components/PortfolioHoldings.jsx`", "Holding entry, instrument search, refresh action, FX-aware summaries, sector/country views, and snapshot presentation"),
            ("`src/lib/market.js`", "Frontend API client for market search, refresh, FX fetch, and browser-side rescue fallbacks for some symbols"),
            ("`server/market-runtime.mjs`", "Search holdings, refresh equities/crypto/commodities, use official NSE EOD for Indian symbols, and provide FX data"),
            ("`api/market/*.js`", "Vercel production handlers for market search, refresh, and FX"),
        ],
        widths=[Inches(2.2), Inches(5.1)],
    )
    add_bullets(doc, [
        "Indian stocks and ETFs now use official NSE end-of-day pricing as the most trustworthy free path for daily refresh.",
        "Mutual funds use a free AMFI-backed NAV feed, which is appropriate for daily refresh instead of intraday pricing.",
        "Global stocks use provider fallbacks such as Twelve Data, Finnhub, Alpha Vantage, and provider-aware symbol inference; crypto and commodities have their own fallbacks.",
        "Holdings store units, average cost, account/broker context, quote symbol overrides, currency, latest fetched price, current value, invested value, unrealized P&L, and planner-oriented metadata such as SIP or realized gain notes.",
        "Daily snapshot history is stored with the user profile so portfolio value survives restarts and cross-device access.",
    ])

    doc.add_heading("10. Local Development Workflow", level=1)
    add_paragraph(doc, "Local development ended up with two execution modes: frontend-only via Vite, and combined app-plus-proxy via a helper script. The combined workflow is preferred because AI routes, admin-aware health checks, market refresh, and FX all depend on the local proxy.")
    add_table(
        doc,
        ["Command", "Purpose"],
        [
            ("`npm run dev`", "Start the Vite frontend only"),
            ("`npm run proxy`", "Start the local AI proxy at 127.0.0.1:8787"),
            ("`npm run dev:full`", "Start both Vite and the AI/market proxy together, reusing an existing proxy if already running"),
            ("`npm run build`", "Create a production build and validate bundling"),
            ("`npm run admin:claim -- --email <user> --project <id>`", "One-time script to assign Firebase auth custom claims such as `role=admin`"),
        ],
        widths=[Inches(1.8), Inches(5.5)],
    )
    add_paragraph(doc, "The Vite dev server proxies `/api/*` to the local AI proxy, so local AI or market failures often come down to whether port 8787 is listening. Once that was understood, the stable local workflow became running both services together rather than starting them separately and accidentally losing the proxy. The app is now pinned to a predictable localhost dev port to avoid earlier port-hopping confusion.")

    doc.add_heading("11. Production Deployment on Vercel", level=1)
    add_paragraph(doc, "The frontend deploys to Vercel as a static React application. The backend runtimes deploy to Vercel by placing handlers in the `/api/ai` and `/api/market` directories. The app also ships a manifest and service worker so the deployed HTTPS build can be installed as a PWA. This is critical: the local Node proxy used in development is not available automatically in production.")
    add_bullets(doc, [
        "The app frontend is built by Vite and served by Vercel.",
        "The `/api/ai/health`, `/api/ai/insights`, and `/api/ai/query` routes are deployed as Vercel functions.",
        "The `/api/market/search`, `/api/market/refresh`, and `/api/market/fx` routes are deployed as Vercel functions for portfolio search, refresh, and currency conversion.",
        "The Vercel project stores runtime secrets as environment variables, then the app is redeployed so the new variables are visible to the functions.",
        "Production AI failures were solved by converting handlers to Node-style `export default async function handler(req, res)` functions and redeploying.",
        "A `vercel.json` layer was added so PWA assets such as `sw.js` and `manifest.webmanifest` are delivered with the right cache behavior.",
    ])
    add_paragraph(doc, "Deployment also required standard Git hygiene. Secrets accidentally committed to `.env` triggered GitHub push protection. The durable fix was to add `.env` to `.gitignore`, remove it from Git tracking, amend the offending commit, and rotate exposed provider keys. Firebase Authentication authorized domains also need to include the deployed Vercel hostname so Google and magic-link sign-in work in preview and production.")

    doc.add_heading("12. Environment Variables and Secrets", level=1)
    add_table(
        doc,
        ["Variable", "Used for", "Where set"],
        [
            ("AI_PROXY_HOST", "Local proxy host", "Local `.env`"),
            ("AI_PROXY_PORT", "Local proxy port", "Local `.env`"),
            ("OPENAI_API_KEY", "OpenAI provider access", "Local `.env` and Vercel env"),
            ("OPENAI_MODEL", "OpenAI model selection", "Local `.env` and Vercel env"),
            ("OPENROUTER_API_KEY", "OpenRouter provider access", "Local `.env` and Vercel env"),
            ("OPENROUTER_MODEL", "OpenRouter model selection", "Local `.env` and Vercel env"),
            ("ANTHROPIC_API_KEY", "Anthropic provider access", "Local `.env` and Vercel env"),
            ("ANTHROPIC_MODEL", "Anthropic model selection", "Local `.env` and Vercel env"),
            ("TWELVE_DATA_API_KEY", "Global equity and crypto provider access", "Local `.env` and Vercel env"),
            ("FINNHUB_API_KEY", "Global equity and crypto fallback access", "Local `.env` and Vercel env"),
            ("ALPHA_VANTAGE_API_KEY", "Stock symbol search and daily price refresh", "Local `.env` and Vercel env"),
            ("FINWISE_ADMIN_EMAILS", "Backend allowlist bridge for admin UI and health access", "Local `.env` and Vercel env"),
        ],
        widths=[Inches(2.2), Inches(2.4), Inches(2.2)],
    )
    add_paragraph(doc, "Useful guardrails were added to backend health reporting: if an OpenAI-style key appears to be stored under `OPENROUTER_API_KEY`, the health payload warns about it; if the viewer is not an admin, health responses return only a safe public shape. This helped catch provider mismatches while keeping sensitive backend details out of everyday user settings.")

    doc.add_heading("13. Security and Reliability Decisions", level=1)
    add_bullets(doc, [
        "Provider keys were moved out of the browser and into backend env vars.",
        "Free market-data usage is limited to manual refresh and a small daily volume, which keeps stock pricing practical without paid infrastructure.",
        "Email-link auth replaced more fragile phone and Apple experiments for this product context.",
        "Git tracking of `.env` was removed to satisfy push protection and reduce accidental exposure.",
        "The app falls back to local heuristics when provider-backed AI is unavailable, so users still receive useful financial guidance.",
        "The account area includes verification, password reset, and email change to make identity management feel complete rather than bolted on.",
        "User-owned planning state now persists in Firebase instead of depending only on browser-local storage, so it survives dev-server restarts and cross-device sign-in.",
        "Admin-only settings are now gated by Firebase claims or a backend email allowlist, and backend health responses also respect that role distinction.",
        "Vault uploads are local-first: files appear immediately, then sync to Firebase Storage in the background with explicit sync-status indicators and retry support.",
    ])

    doc.add_heading("14. Performance Optimization and Code Splitting", level=1)
    add_paragraph(doc, "Once the live app was stable, the next improvement was performance. The app had grown large enough that a single big production bundle was no longer ideal. The final step was to lazy-load feature screens and split vendor chunks by responsibility.")
    add_table(
        doc,
        ["Optimization", "Effect"],
        [
            ("Lazy-loaded tab components", "Users only download analytics, AI, goals, import, or net-worth code when those areas are opened"),
            ("Manual vendor chunking in Vite", "React, Firebase, and chart libraries are split into separate chunks"),
            ("Small tab loading state", "Transitions feel intentional while a deferred screen chunk loads"),
            ("PWA shell caching", "The service worker can keep the shell feeling faster and more app-like on repeat visits"),
        ],
        widths=[Inches(2.3), Inches(5.2)],
    )
    add_paragraph(doc, "After optimization, the main app shell bundle dropped into a much healthier range, with charting and Firebase each living in their own chunks instead of being paid on first paint.")

    doc.add_heading("15. Implementation Journey and Major Decisions", level=1)
    add_image_center(doc, FLOW_IMG, 6.9)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Figure 2. Implementation and delivery lifecycle")
    run.italic = True
    run.font.color.rgb = MUTED
    add_paragraph(doc, "The project evolved in layers rather than one rewrite. First the finance surface area expanded. Then the giant root component was split into modular screens. Then AI security moved server-side. Then the auth model was simplified. Finally, the UX and performance passes tightened the product for real-world use.")

    doc.add_heading("16. Testing and Validation Checklist", level=1)
    add_bullets(doc, [
        "Build verification with `npm run build` after each meaningful implementation phase",
        "Health checks against local AI proxy and Vite-proxied `/api/ai/health`",
        "Market search and refresh verification against local `/api/market/*` routes",
        "Official NSE EOD fallback verification for Indian stock symbols",
        "End-to-end AI query checks against configured providers",
        "Authentication checks for email/password, Google, magic-link completion, password reset, and verification flows",
        "Role-aware admin checks for Settings and backend health responses",
        "PWA install-path checks on deployed HTTPS URLs",
        "Vault upload, background sync, retry, export/import, and migration checks",
        "Visual refinement of AI layouts, dashboard diagnostics, and settings/account flows",
        "Production validation on Vercel after env changes and handler updates",
    ])

    doc.add_heading("17. Troubleshooting Reference", level=1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Recommended fix"],
        [
            ("`ECONNREFUSED 127.0.0.1:8787` in local AI requests", "Local proxy not running", "Use `npm run dev:full` or start `npm run proxy` separately"),
            ("AI falls back to local reasoning in Vercel", "No deployed `/api` runtime, missing env vars, or no redeploy", "Ensure `api/ai/*` exists, env vars are set in Vercel, and trigger a new deployment"),
            ("Market holdings vanish after a server restart", "State only existed in browser storage or pre-migration local data", "Persist holdings in Firebase-backed `meta/appState` and let the first authenticated load migrate local fallback state"),
            ("Stock refresh/search fails", "Missing provider key, stale symbol mapping, or provider coverage issue", "Check provider env vars, use quote symbol overrides, and rely on official NSE EOD for Indian holdings"),
            ("Commodity refresh fails", "Provider path returned no usable quote", "Use the commodity symbol mappings and Yahoo chart fallback now built into the market runtime"),
            ("Magic link shows success but email appears missing", "Delivery landed in Promotions or Spam", "Check inbox tabs and spam folder, then brand templates later if needed"),
            ("Phone test numbers work but real OTP fails", "Firebase production SMS constraints and billing/setup complexity", "Remove or deprioritize phone OTP for this product"),
            ("GitHub push rejected", "Secret detected in commit history", "Remove `.env` from tracking, amend commit, rotate keys, push again"),
            ("OpenRouter or OpenAI requests fail", "Wrong key placement, quota problem, or unsupported model", "Verify provider key name, billing/quota, and model selection"),
            ("Blank page on localhost after a code change", "Frontend runtime error in the app shell", "Check Vite console/runtime logs; in this project one known example was using `hasPendingVaultSync` before declaration in `App.jsx`"),
            ("Admin section still shows `USER`", "Proxy not restarted after env change, or no Firebase claim", "Restart proxy, sign out/in, or use the one-time admin claim script"),
        ],
        widths=[Inches(2.2), Inches(2.3), Inches(2.9)],
    )

    doc.add_heading("18. File and Folder Reference", level=1)
    add_table(
        doc,
        ["Path", "Role"],
        [
            ("`src/App.jsx`", "Stateful app shell, routing by tab, derived metrics, auth orchestration"),
            ("`src/components/*`", "Feature modules: dashboard, AI, analytics, goals, net worth, import, history, settings, login"),
            ("`src/components/PortfolioHoldings.jsx`", "Dedicated holdings entry, price refresh, and portfolio summary surface"),
            ("`src/components/DocumentsVault.jsx`", "Document metadata, local-first uploads, sync status, and vault controls"),
            ("`src/components/TimelineView.jsx`", "Timeline/calendar-oriented reminders and milestone view"),
            ("`src/lib/utils.js`", "Formatting helpers, date logic, pie aggregation, storage keys"),
            ("`src/lib/ai.js`", "Heuristic report generation and frontend AI request helpers"),
            ("`src/lib/market.js`", "Frontend market search and refresh API helpers"),
            ("`src/lib/passkey.js`", "Local passkey helpers for supported secure environments"),
            ("`src/lib/vaultStorage.js`", "IndexedDB staging, Firebase Storage sync, vault migration/import/export helpers"),
            ("`src/lib/i18n.js`", "Language dictionary and translation lookup"),
            ("`src/styles/appStyles.js`", "Global visual system"),
            ("`src/firebase.js`", "Firebase bootstrapping"),
            ("`server/ai-runtime.mjs`", "Shared provider runtime and health logic"),
            ("`server/market-runtime.mjs`", "Shared market search and refresh runtime"),
            ("`server/ai-proxy.mjs`", "Local development AI proxy"),
            ("`scripts/dev-full.mjs`", "Combined app/proxy development launcher with proxy reuse"),
            ("`scripts/set-admin-claim.mjs`", "One-time Firebase admin-claim assignment helper"),
            ("`api/ai/*.js`", "Production serverless AI handlers for Vercel"),
            ("`api/market/*.js`", "Production serverless market-data handlers for Vercel"),
            ("`public/manifest.webmanifest` + `public/sw.js`", "PWA metadata and service worker"),
            ("`vercel.json`", "Vercel routing/cache behavior for PWA and serverless deployment"),
            ("`vite.config.js`", "Dev proxy rules, fixed dev port behavior, and vendor chunk splitting"),
        ],
        widths=[Inches(2.3), Inches(5.2)],
    )

    doc.add_heading("19. Recommended Next Steps", level=1)
    add_bullets(doc, [
        "Brand the Firebase email templates so inbox trust is stronger and magic-link delivery feels more polished.",
        "Introduce OCR and AI Q&A over uploaded vault documents once the storage layer is fully settled.",
        "Add richer onboarding data seeding for first-time users importing from bank statements or choosing demo mode.",
        "Track release metrics such as AI usage, import completion rate, and goal creation rate after launch.",
        "Extend market coverage with richer symbol hints, portfolio performance attribution, and a more deliberate commodity/global quote strategy once the core daily-refresh model is stable.",
        "Promote the install/PWA flow with a more polished in-app install banner and offline-ready indicator.",
    ])

    appendix = doc.add_section(WD_SECTION.NEW_PAGE)
    add_page_number(appendix)
    doc.add_heading("Appendix: Deployment Runbook", level=1)
    add_bullets(doc, [
        "1. Confirm Firebase Authentication methods are enabled: Email/Password, Google, and Email Link.",
        "2. Confirm Firestore rules and project configuration are correct.",
        "3. Add provider keys in Vercel Project Settings -> Environment Variables.",
        "4. Add `ALPHA_VANTAGE_API_KEY`, `TWELVE_DATA_API_KEY`, `FINNHUB_API_KEY`, and `FINWISE_ADMIN_EMAILS` if those production capabilities should work.",
        "5. Add deployed Vercel domains to Firebase Authentication authorized domains.",
        "6. Ensure Firebase Storage rules permit authenticated user uploads and reads for vault files.",
        "7. Redeploy after env changes.",
        "8. Validate `/api/ai/health`, `/api/market/search`, and `/api/market/fx` in production.",
        "9. Test sign-in, AI query, transaction add, import, analytics, goals, net worth, holdings refresh, vault upload/sync, install prompt behavior, and persistence reload flows.",
        "10. Check browser console and Vercel function logs for any last-mile failures.",
    ])

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_DOCX)
